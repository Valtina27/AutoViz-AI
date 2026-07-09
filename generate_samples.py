"""
Sample data generator for testing the AI Data Visualization application
Creates sample Excel, Word, and Text files with various data patterns
"""

import pandas as pd
from docx import Document
import random
from datetime import datetime, timedelta

def generate_sample_excel():
    """Generate a sample Excel file with multiple sheets"""
    
    # Sheet 1: Product Sales Data
    products = ['Phone', 'Laptop', 'Watch', 'Tablet', 'Headset']
    months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun']
    
    sales_data = {
        'Product': products * len(months),
        'Month': [m for m in months for _ in products],
        'Sales': [random.randint(50, 200) for _ in range(len(products) * len(months))],
        'Revenue': [random.randint(10000, 50000) for _ in range(len(products) * len(months))]
    }
    
    # Sheet 2: Employee Performance
    employees = ['Alice', 'Bob', 'Charlie', 'Diana', 'Eve']
    departments = ['Sales', 'Marketing', 'IT', 'HR', 'Finance']
    
    performance_data = {
        'Employee': employees,
        'Department': random.sample(departments, len(employees)),
        'Performance_Score': [random.randint(60, 100) for _ in employees],
        'Projects_Completed': [random.randint(5, 20) for _ in employees],
        'Years_Experience': [random.randint(1, 10) for _ in employees]
    }
    
    # Create Excel file with multiple sheets
    with pd.ExcelWriter('sample_data.xlsx', engine='openpyxl') as writer:
        pd.DataFrame(sales_data).to_excel(writer, sheet_name='Sales', index=False)
        pd.DataFrame(performance_data).to_excel(writer, sheet_name='Performance', index=False)
    
    print("✓ Created: sample_data.xlsx (2 sheets)")

def generate_sample_word():
    """Generate a sample Word document with text and tables"""
    
    doc = Document()
    
    # Add title
    doc.add_heading('Quarterly Business Report', 0)
    
    # Add some text
    doc.add_paragraph('This report summarizes our business performance for Q4 2024.')
    doc.add_paragraph('Key Metrics: Revenue: 450000, Customers: 1250, Growth Rate: 15.5%')
    
    # Add a table
    doc.add_heading('Regional Sales Data', level=1)
    
    regions = ['North', 'South', 'East', 'West', 'Central']
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Region'
    hdr_cells[1].text = 'Sales'
    hdr_cells[2].text = 'Target'
    
    # Data rows
    for region in regions:
        row_cells = table.add_row().cells
        row_cells[0].text = region
        row_cells[1].text = str(random.randint(50000, 150000))
        row_cells[2].text = str(random.randint(60000, 140000))
    
    # Add more text
    doc.add_paragraph()
    doc.add_paragraph('Customer Satisfaction: 4.5/5.0, Employee Count: 85, New Hires: 12')
    
    doc.save('sample_report.docx')
    print("✓ Created: sample_report.docx (with tables)")

def generate_sample_text():
    """Generate sample text files"""
    
    # CSV file
    csv_data = """Category,Value,Date,Status
Electronics,125,2024-01-15,Active
Furniture,89,2024-02-20,Active
Clothing,156,2024-03-10,Completed
Electronics,143,2024-04-05,Active
Furniture,67,2024-05-12,Pending
Clothing,201,2024-06-18,Completed
Electronics,178,2024-07-22,Active
"""
    
    with open('sample_data.csv', 'w') as f:
        f.write(csv_data)
    
    print("✓ Created: sample_data.csv")
    
    # Text file with structured data
    text_data = """Product Analysis Report
    
Summary Statistics:
Total Products: 150
Average Price: 249.99
Top Category: Electronics
Market Share: 32%

Performance Metrics:
Q1 Revenue: 350000
Q2 Revenue: 425000
Q3 Revenue: 480000
Q4 Revenue: 510000

Customer Insights:
New Customers: 450
Returning Customers: 1200
Satisfaction Rate: 87%
"""
    
    with open('sample_report.txt', 'w') as f:
        f.write(text_data)
    
    print("✓ Created: sample_report.txt")

def generate_all_samples():
    """Generate all sample files"""
    
    print("\nGenerating sample data files...")
    print("-" * 50)
    
    try:
        generate_sample_excel()
        generate_sample_word()
        generate_sample_text()
        
        print("-" * 50)
        print("\n✓ All sample files created successfully!")
        print("\nYou can now test the application by uploading these files:")
        print("  - sample_data.xlsx")
        print("  - sample_report.docx")
        print("  - sample_data.csv")
        print("  - sample_report.txt")
        print("\nStart the app with: python app.py")
        
    except Exception as e:
        print(f"\n✗ Error generating sample files: {e}")

if __name__ == "__main__":
    generate_all_samples()
